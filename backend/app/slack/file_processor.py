"""File processing module for parsing uploaded performance data files."""

import io
import json
from typing import Any, Dict, List, Optional
from pathlib import Path

import pandas as pd
from docx import Document
from pypdf import PdfReader
import structlog

logger = structlog.get_logger(__name__)


class FileProcessor:
    """Process various file formats to extract performance data and context."""

    SUPPORTED_EXTENSIONS = {".csv", ".xlsx", ".xls", ".pdf", ".docx", ".json"}

    def __init__(self):
        """Initialize the file processor."""
        self._processors = {
            ".csv": self._process_csv,
            ".xlsx": self._process_excel,
            ".xls": self._process_excel,
            ".pdf": self._process_pdf,
            ".docx": self._process_docx,
            ".json": self._process_json,
        }

    def can_process(self, filename: str) -> bool:
        """
        Check if a file can be processed.

        Args:
            filename: Name of the file.

        Returns:
            True if the file type is supported.
        """
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS

    async def process_file(
        self,
        file_content: bytes,
        filename: str,
    ) -> Dict[str, Any]:
        """
        Process an uploaded file and extract relevant data.

        Args:
            file_content: Raw bytes of the file.
            filename: Name of the file for determining type.

        Returns:
            Dictionary with extracted data and metadata.
        """
        ext = Path(filename).suffix.lower()

        if ext not in self._processors:
            raise ValueError(f"Unsupported file type: {ext}")

        logger.info("processing_file", filename=filename, extension=ext)

        processor = self._processors[ext]
        result = await processor(file_content, filename)

        logger.info("file_processed", filename=filename, data_type=result.get("type"))

        return result

    async def _process_csv(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process a CSV file."""
        try:
            df = pd.read_csv(io.BytesIO(content))
            return self._dataframe_to_result(df, filename, "csv")
        except Exception as e:
            logger.error("csv_processing_error", filename=filename, error=str(e))
            return {
                "type": "error",
                "filename": filename,
                "error": str(e),
            }

    async def _process_excel(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process an Excel file."""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(io.BytesIO(content))
            sheets = {}

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                sheets[sheet_name] = df.to_dict(orient="records")

            return {
                "type": "spreadsheet",
                "filename": filename,
                "sheets": sheets,
                "sheet_names": excel_file.sheet_names,
                "summary": self._generate_spreadsheet_summary(sheets),
            }
        except Exception as e:
            logger.error("excel_processing_error", filename=filename, error=str(e))
            return {
                "type": "error",
                "filename": filename,
                "error": str(e),
            }

    async def _process_pdf(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process a PDF file."""
        try:
            reader = PdfReader(io.BytesIO(content))
            text_content = []

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_content.append({
                        "page": page_num + 1,
                        "content": text.strip(),
                    })

            full_text = "\n\n".join([p["content"] for p in text_content])

            return {
                "type": "document",
                "filename": filename,
                "format": "pdf",
                "page_count": len(reader.pages),
                "text_content": full_text,
                "pages": text_content,
            }
        except Exception as e:
            logger.error("pdf_processing_error", filename=filename, error=str(e))
            return {
                "type": "error",
                "filename": filename,
                "error": str(e),
            }

    async def _process_docx(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process a DOCX file."""
        try:
            doc = Document(io.BytesIO(content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract tables if present
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)

            return {
                "type": "document",
                "filename": filename,
                "format": "docx",
                "text_content": "\n\n".join(paragraphs),
                "paragraph_count": len(paragraphs),
                "tables": tables,
                "table_count": len(tables),
            }
        except Exception as e:
            logger.error("docx_processing_error", filename=filename, error=str(e))
            return {
                "type": "error",
                "filename": filename,
                "error": str(e),
            }

    async def _process_json(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process a JSON file."""
        try:
            data = json.loads(content.decode("utf-8"))

            return {
                "type": "json",
                "filename": filename,
                "data": data,
                "summary": self._generate_json_summary(data),
            }
        except Exception as e:
            logger.error("json_processing_error", filename=filename, error=str(e))
            return {
                "type": "error",
                "filename": filename,
                "error": str(e),
            }

    def _dataframe_to_result(
        self,
        df: pd.DataFrame,
        filename: str,
        format_type: str,
    ) -> Dict[str, Any]:
        """Convert a DataFrame to a result dictionary."""
        # Detect if this looks like performance data
        performance_columns = {
            "spend", "cost", "budget", "impressions", "clicks",
            "conversions", "cpa", "cpc", "roas", "revenue",
            "campaign", "ad_set", "adset", "ad_group", "platform"
        }

        columns_lower = {col.lower().replace(" ", "_") for col in df.columns}
        is_performance_data = bool(columns_lower & performance_columns)

        return {
            "type": "performance_data" if is_performance_data else "tabular",
            "filename": filename,
            "format": format_type,
            "columns": list(df.columns),
            "row_count": len(df),
            "data": df.to_dict(orient="records"),
            "summary": self._generate_dataframe_summary(df),
            "is_performance_data": is_performance_data,
        }

    def _generate_dataframe_summary(self, df: pd.DataFrame) -> str:
        """Generate a text summary of a DataFrame."""
        lines = [
            f"Rows: {len(df)}",
            f"Columns: {', '.join(df.columns)}",
        ]

        # Add numeric column statistics
        numeric_cols = df.select_dtypes(include=["number"]).columns
        if len(numeric_cols) > 0:
            lines.append("\nNumeric Summary:")
            for col in numeric_cols[:5]:  # Limit to first 5
                lines.append(f"  {col}: min={df[col].min():.2f}, max={df[col].max():.2f}, mean={df[col].mean():.2f}")

        return "\n".join(lines)

    def _generate_spreadsheet_summary(self, sheets: Dict[str, List]) -> str:
        """Generate a summary for a multi-sheet spreadsheet."""
        lines = [f"Sheets: {len(sheets)}"]
        for name, data in sheets.items():
            lines.append(f"  - {name}: {len(data)} rows")
        return "\n".join(lines)

    def _generate_json_summary(self, data: Any) -> str:
        """Generate a summary for JSON data."""
        if isinstance(data, list):
            return f"Array with {len(data)} items"
        elif isinstance(data, dict):
            return f"Object with keys: {', '.join(list(data.keys())[:10])}"
        else:
            return f"Value of type {type(data).__name__}"
